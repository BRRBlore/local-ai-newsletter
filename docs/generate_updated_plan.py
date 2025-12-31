import os
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    # Just in case
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # --- STYLES ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- TITLE PAGE ---
    title = doc.add_heading('AI Newsletter Project: Master Plan v2.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('The "Hardened" Hybrid Architecture (SQLite + Ollama + Vector Ranking)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph('\n\n')
    doc.add_paragraph('Status: Finalized for Implementation')
    doc.add_paragraph('Architecture: Option 2 (Hybrid)')
    doc.add_paragraph('Key Upgrade: "Pro" Database Schema & Robust ETL')
    doc.add_page_break()

    # --- EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This document outlines the implementation of a production-grade automated AI newsletter. "
        "We have selected the 'Hybrid Option' which balances Data Science learning (Vector Clustering) "
        "with robust software engineering (SQLite History, YAML Config, Env Secrets)."
    )

    # --- ARCHITECTURE ---
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph("The pipeline runs daily in 5 distinct stages:")
    
    diagram = """
    [ INTERNET ]      [ LOCAL MACHINE ]                     [ SUBSCRIBERS ]
         |
    (1. ETL Fetch) -> [ src/fetcher.py ] 
         |            (Parses RSS -> Normalizes -> Hashes)
         |                    |
         |            [ data/newsletter.db ] (Raw 'items' table)
         |                    |
         |            (2. AI Logic)
         |            [ src/ranker.py ]
         |            (Embeddings -> Clustering -> Scoring)
         |                    |
         |            (3. Generation)
         |            [ src/editor.py ] (Ollama LLM)
         |            (Summarizes Top 20 -> JSON)
         |                    |
         |            (4. Delivery)
         |            [ src/sender.py ] (SMTP) -----------> [ Email Inbox ]
    """
    p = doc.add_paragraph(diagram)
    try:
        p.style = 'Quote' 
    except:
        pass # Fallback if style missing

    # --- PHASE 1: FOUNDATION ---
    doc.add_heading('3. Phase 1: The "Hardened" Foundation', level=1)
    doc.add_paragraph(
        "We prioritize data integrity and security from Day 1. "
        "This setup prevents technical debt."
    )

    doc.add_heading('3.1 Project Structure', level=2)
    structure = """
    ai_newsletter/
     ├── .env                  # Secrets (API Keys, Passwords) - NEVER COMMIT
     ├── config/
     │   └── sources.yaml      # List of 20+ feeds with Trust Scores
     ├── data/
     │   └── newsletter.db     # SQLite Database
     ├── src/
     │   ├── database.py       # DB Connection & Schema Init
     │   ├── fetcher.py        # ETL Logic
     │   ├── ranker.py         # AI Clustering Logic
     │   └── sender.py         # Email Logic
     └── main.py               # Daily Orchestrator
    """
    p = doc.add_paragraph(structure)
    try: p.style = 'Quote' 
    except: pass

    doc.add_heading('3.2 The "Pro" Database Schema', level=2)
    doc.add_paragraph(
        "We use a normalized schema to separate 'Raw Data' from 'derived insights'."
    )
    sql_code = """
    -- 1. RAW ITEMS (The Data Lake)
    CREATE TABLE items (
        id INTEGER PRIMARY KEY,
        url TEXT UNIQUE,
        title TEXT,
        source_name TEXT,
        published_at TIMESTAMP,
        content_hash TEXT  -- For deduplication
    );

    -- 2. CLUSTERS (The AI Features)
    CREATE TABLE clusters (
        cluster_id TEXT PRIMARY KEY,
        run_date DATE,
        representative_item_id INTEGER,
        size INTEGER,
        score REAL
    );

    -- 3. HISTORY (The Memory)
    CREATE TABLE seen_items (
        canonical_url TEXT PRIMARY KEY,
        last_sent_at TIMESTAMP
    );
    """
    p = doc.add_paragraph(sql_code)
    try: p.style = 'Quote' 
    except: pass

    # --- PHASE 2: CONFIG ---
    doc.add_heading('4. Phase 2: Configuration & Secrets', level=1)
    doc.add_paragraph("We avoid hardcoding. All logic is driven by config files.")
    
    doc.add_heading('4.1 sources.yaml Strategy', level=2)
    yaml_ex = """
    sources:
      - name: "The Rundown"
        url: "https://www.therundown.ai/feed"
        trust_score: 10
        type: "rss"
        max_items: 10
        enabled: true
    
      - name: "Ben's Bites"
        url: "https://bensbites.beehiiv.com/feed"
        trust_score: 9
        enabled: true
    """
    p = doc.add_paragraph(yaml_ex)
    try: p.style = 'Quote' 
    except: pass

    # --- PHASE 3: AI LOGIC ---
    doc.add_heading('5. Phase 3: The "Buzz" Algorithm', level=1)
    doc.add_paragraph(
        "How we determine 'What's Hot' without human intervention."
    )
    doc.add_heading('5.1 The Scoring Formula', level=2)
    doc.add_paragraph(
        "Score = (Cluster_Size * 10) + (Source_Trust * 5) + (Recency_Decay) - (Duplicate_Penalty)"
    )
    doc.add_paragraph(
        "Where 'Recency_Decay' reduces the score if the news is > 24 hours old."
    )

    # --- APPENDIX ---
    doc.add_heading('Appendix: The Setup Script', level=1)
    doc.add_paragraph("Run this code to initialize the Version 2.0 folder structure immediately.")
    
    setup_code = """
import os, sqlite3

def setup_v2():
    # 1. Folders
    for f in ['config', 'data', 'src', 'logs']:
        os.makedirs(f, exist_ok=True)
    
    # 2. Database
    conn = sqlite3.connect('data/newsletter.db')
    c = conn.cursor()
    
    # Create ITEMS table
    c.execute('''CREATE TABLE IF NOT EXISTS items 
                 (id INTEGER PRIMARY KEY, url TEXT UNIQUE, title TEXT, 
                  source_name TEXT, published_at TIMESTAMP)''')
                  
    # Create SUBSCRIBERS table
    c.execute('''CREATE TABLE IF NOT EXISTS subscribers 
                 (email TEXT UNIQUE, name TEXT, is_active BOOLEAN)''')
                 
    conn.commit()
    conn.close()
    
    # 3. Config
    with open('config/sources.yaml', 'w') as f:
        f.write('sources:\\n  - name: Example\\n    url: http://example.com/rss')
        
    print("✅ Version 2.0 Environment Ready.")

if __name__ == "__main__":
    setup_v2()
    """
    p = doc.add_paragraph(setup_code)
    try: p.style = 'Quote' 
    except: pass

    # Save
    doc.save('AI_Newsletter_Plan_v2.docx')
    print("✅ Success! 'AI_Newsletter_Plan_v2.docx' created.")

if __name__ == "__main__":
    create_document()