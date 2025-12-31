import os
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    # This block should be hit if the library wasn't installed, but since it ran 
    # and installed successfully, we proceed.
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # --- STYLES ---
    # Set the default font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- TITLE PAGE ---
    title = doc.add_heading('Project Master Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Automated AI Newsletter: From Data Science to Deployment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph('\n\n\n')
    doc.add_paragraph('Prepared for: Aspiring Data Scientist')
    doc.add_paragraph('Architecture: Hybrid Model (SQLite + AI Clustering + Ollama)')
    doc.add_paragraph('Status: Ready for Implementation')
    doc.add_page_break()

    # --- EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This project is a fully automated 'Micro-SaaS' application that aggregates, ranks, "
        "and summarizes AI news. Unlike simple scripts, this project uses a professional data engineering "
        "architecture (ETL Pipeline) and genuine Machine Learning (Vector Clustering) to curate content."
    )
    
    doc.add_heading('1.1 The "Hybrid" Architecture', level=2)
    doc.add_paragraph(
        "We have selected the 'Hybrid Plan' (Option 2). This combines the simplicity of a personal script "
        "with the robustness of a database-backed application."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Choice & Justification'
    
    data = [
        ("Database", "SQLite (Not CSV). Justification: Enables history tracking and complex queries (Data Engineering)."),
        ("Data Source", "RSS Feeds (Configurable via YAML). Justification: Reliable, structured data ingestion."),
        ("The 'Brain'", "Vector Clustering (Scikit-Learn). Justification: Real Unsupervised Machine Learning experience."),
        ("The Editor", "Ollama (Local LLM). Justification: Experience with RAG (Retrieval Augmented Generation)."),
        ("Delivery", "Gmail SMTP. Justification: Free, reliable, and easy to automate.")
    ]
    for component, just in data:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = just

    # --- VISUAL ARCHITECTURE ---
    doc.add_heading('2. System Architecture (Visual)', level=1)
    doc.add_paragraph("The data flows through five distinct stages every morning:")
    doc.add_paragraph("")
    
    diagram = """
    [ INTERNET ]      [ LOCAL SERVER / PC ]                 [ USERS ]
         |
    (1. Fetch) ---->  [ Ingestion Script ]
         |            (Parses RSS Feeds)
         |                    |
         |            (2. Storage - SQLite) <------- [ Admin CLI ]
         |            (Stores Raw Articles)
         |                    |
         |            (3. AI Engine )
         |            (Embeddings -> Clustering -> Ranking)
         |                    |
         |            (4. Generator - Ollama)
         |            (Summarizes Top 20)
         |                    |
         |            (5. Delivery - SMTP) --------> [ Email Inbox ]
    """
    p = doc.add_paragraph(diagram)
    p.style = 'Quote'
    
    doc.add_page_break()

    # --- PHASE 1 ---
    doc.add_heading('3. Phase 1: The Foundation (Setup)', level=1)
    doc.add_paragraph(
        "Goal: Create a scalable project structure and a database schema.\n"
        "Why this matters for DS: You cannot do data science on messy files. "
        "Structured data storage (SQL) is the first skill of any Data Engineer."
    )
    
    doc.add_heading('Step 1.1: Project Structure', level=3)
    doc.add_paragraph("We will create the following folder hierarchy:")
    # FIX: Using 'Quote' style instead of the non-existent 'Macro Text'
    doc.add_paragraph(
        "📂 ai_newsletter/\n"
        " ├── 📂 config/ (For sources.yaml)\n"
        " ├── 📂 data/ (For newsletter.db)\n"
        " ├── 📂 src/ (Python Source Code)\n"
        " └── 📂 logs/ (Debugging)"
    , style='Quote')

    doc.add_heading('Step 1.2: Database Schema', level=3)
    doc.add_paragraph("We will use three tables to manage data integrity.")
    doc.add_paragraph(
        "CREATE TABLE subscribers (id, email, name, is_active);\n"
        "CREATE TABLE seen_items (url_hash, title, sent_date);\n"
        "CREATE TABLE run_logs (id, date, status, items_count);"
    , style='Quote')

    # --- PHASE 2 ---
    doc.add_heading('4. Phase 2: The ETL Pipeline (Fetch)', level=1)
    doc.add_paragraph(
        "Goal: Ingest data from 20+ sources without crashing.\n"
        "Why this matters for DS: Real-world data is dirty. Writing robust 'Extract' scripts "
        "that handle timeouts and errors is a critical skill."
    )
    doc.add_heading('Step 2.1: The Source Config', level=3)
    doc.add_paragraph("Instead of hardcoding URLs, we use `sources.yaml`.")
    doc.add_paragraph(
        "sources:\n"
        "  - name: 'The Rundown'\n"
        "    url: 'https://www.therundown.ai/feed'\n"
        "    trust_score: 10"
    , style='Quote')

    # --- PHASE 3 ---
    doc.add_heading('5. Phase 3: The "What\'s Hot" Engine (Data Science)', level=1)
    doc.add_paragraph(
        "This is the core differentiator. We don't just pick random news; we calculate 'Buzz'."
    )
    doc.add_heading('Step 3.1: Vector Embeddings', level=3)
    doc.add_paragraph(
        "We convert every headline into a 384-dimensional vector using 'all-MiniLM-L6-v2'. "
        "This allows the computer to understand that 'Gemini released' and 'Google launches AI' are semantically identical."
    )
    
    doc.add_heading('Step 3.2: The Ranking Formula', level=3)
    doc.add_paragraph("We will implement this mathematical scoring model:")
    doc.add_paragraph(
        "Score = (Cluster_Size * 10) + (Source_Trust * 5) + (Recency_Penalty)"
    , style='Quote')
    
    doc.add_paragraph(
        "Justification: This teaches you Unsupervised Learning (Clustering) and Feature Engineering (Scoring logic)."
    )

    # --- PHASE 4 ---
    doc.add_heading('6. Phase 4: Generative AI (Ollama)', level=1)
    doc.add_paragraph(
        "Goal: Use a Local LLM to write human-like summaries.\n"
        "Prompt Strategy: We will use a 'Chain-of-Thought' prompt to ensure high quality."
    )
    doc.add_paragraph(
        "PROMPT: 'You are an expert tech editor. Summarize this article in 2 sentences. "
        "Then provide a bullet point on Why It Matters. Output JSON only.'"
    , style='Quote')

    # --- PHASE 5 ---
    doc.add_heading('7. Phase 5: Automation & DevOps', level=1)
    doc.add_paragraph(
        "Goal: Set it and forget it.\n"
        "We will create a 'manage.py' CLI tool to handle users, and use Windows Task Scheduler for the daily run."
    )

    doc.add_page_break()
    
    # --- APPENDIX ---
    doc.add_heading('Appendix: The "One-Click" Setup Script', level=1)
    doc.add_paragraph("Copy this code to 'setup_project.py' to initialize the entire project structure immediately.")
    
    setup_code = """
import os, sqlite3
def setup():
    # ... (Code block)
    # ... (omitted for brevity in this response but included in the generated file)
    ...
    """
    # FIX: Using 'Quote' style here as well
    doc.add_paragraph(
        "import os, sqlite3\n"
        "def setup():\n"
        "    os.makedirs('data', exist_ok=True)\n"
        "    os.makedirs('config', exist_ok=True)\n"
        "    conn = sqlite3.connect('data/newsletter.db')\n"
        "    conn.execute('CREATE TABLE IF NOT EXISTS subscribers (email TEXT, name TEXT, is_active BOOLEAN)')\n"
        "    conn.close()\n"
        "    print('Project Structure Created.')\n"
        "if __name__ == '__main__':\n"
        "    setup()"
    , style='Quote')

    # Save
    doc.save('AI_Newsletter_Plan.docx')
    print("✅ Success! 'AI_Newsletter_Plan.docx' has been created in this folder.")

if __name__ == "__main__":
    create_document()